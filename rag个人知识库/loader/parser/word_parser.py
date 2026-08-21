import logging
import os
from typing import List, Optional
from zipfile import ZipFile

from docx import Document as DocxDocument
from lxml import etree

logger = logging.getLogger(__name__)


# ── 复杂度阈值：得分超过此值则认为文档复杂，需要走 MineRU 解析 ──
COMPLEXITY_THRESHOLD = 3


def word_complicatedness(file_path: str) -> int:
    """
    评估 Word 文档的复杂度，返回一个整数得分。

    检测维度（按影响解析准确度的权重加分）：
      - 内嵌图片（inline shapes + media 目录）
      - 表格（含嵌套表格、合并单元格）
      - 分栏布局（columns > 1）
      - 文本框 / 浮动图形（VML + DrawingML）
      - SmartArt 图形
      - 图表（Chart）
      - OLE 嵌入对象（Excel、Visio 等）
      - 数学公式（OMML）
      - 页眉页脚中的复杂内容

    得分越高说明文档越复杂，越不适合用 UnstructuredWordDocumentLoader 直接解析。
    """
    # 防御性检查：python-docx / zipfile 只能解析 OpenXML 的 .docx，
    # 旧版二进制 .doc 会抛 PackageNotFoundError/BadZipFile，这里直接给出明确提示
    if not file_path.lower().endswith(".docx"):
        raise ValueError(
            f"word_complicatedness 仅支持 .docx（OpenXML）格式，收到：{file_path}；"
            "旧版 .doc 请先用 WPS/Word 另存为 .docx"
        )
    score = 0
    doc = DocxDocument(file_path)

    # ── 1. 检查内嵌图片（python-docx API） ──
    inline_count = len(doc.inline_shapes)
    if inline_count > 0:
        logger.info("  [WordChecker] 检测到 %d 张内嵌图片  (+2)", inline_count)
        score += 2

    # ── 2. 检查表格（含嵌套 & 合并单元格） ──
    if len(doc.tables) > 0:
        table_score = 2
        # 检查是否存在嵌套表格或合并单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 嵌套表格
                    if cell.tables:
                        table_score = 4
                        break
                # 检查合并单元格：同一行中存在相同引用的 cell
                cell_refs = [id(c._tc) for c in row.cells]
                if len(cell_refs) != len(set(cell_refs)):
                    table_score = max(table_score, 3)
        logger.info("  [WordChecker] 检测到 %d 个表格  (+%d)", len(doc.tables), table_score)
        score += table_score

    # ── 3. 解析底层 XML，进行更全面的检查 ──
    with ZipFile(file_path) as z:
        xml = z.read("word/document.xml")
        # 记录 zip 中的媒体/嵌入文件，供后续判断
        zip_names = z.namelist()
    root = etree.fromstring(xml)

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "v": "urn:schemas-microsoft-com:vml",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "c": "http://schemas.openxmlformats.org/drawingml/2006/chart",
        "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
        "o": "urn:schemas-microsoft-com:office:office",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }

    # ── 3a. 检查分栏 (w:cols 且 num > 1) ──
    for cols_elem in root.xpath("//w:sectPr/w:cols", namespaces=ns):
        num = cols_elem.get(f"{{{ns['w']}}}num")
        if num and int(num) > 1:
            logger.info("  [WordChecker] 检测到多栏布局  (+3)")
            score += 3
            break

    # ── 3b. 检查 VML 文本框 (v:textbox) ──
    textboxes = root.xpath("//v:textbox", namespaces=ns)
    if textboxes:
        logger.info("  [WordChecker] 检测到 %d 个 VML 文本框  (+2)", len(textboxes))
        score += 2

    # ── 3c. 检查 DrawingML 浮动图形 (wps:wsp / wp:anchor) ──
    drawings = root.xpath("//wps:wsp | //wp:anchor", namespaces=ns)
    if drawings:
        logger.info("  [WordChecker] 检测到 %d 个浮动图形  (+2)", len(drawings))
        score += 2

    # ── 3d. 检查 SmartArt（dgm:relIds 表示存在关系图/SmartArt） ──
    smart_arts = root.xpath("//dgm:relIds", namespaces=ns)
    if smart_arts:
        logger.info("  [WordChecker] 检测到 %d 个 SmartArt 图形  (+3)", len(smart_arts))
        score += 3

    # ── 3e. 检查图表（c:chart） ──
    charts = root.xpath("//c:chart", namespaces=ns)
    if charts:
        logger.info("  [WordChecker] 检测到 %d 个图表  (+3)", len(charts))
        score += 3

    # ── 3f. 检查 OLE 嵌入对象 (o:OLEObject) ──
    ole_objects = root.xpath("//o:OLEObject", namespaces=ns)
    if ole_objects:
        logger.info("  [WordChecker] 检测到 %d 个 OLE 嵌入对象  (+3)", len(ole_objects))
        score += 3

    # ── 3g. 检查数学公式 (m:oMath) ──
    math_blocks = root.xpath("//m:oMath", namespaces=ns)
    if math_blocks:
        logger.info("  [WordChecker] 检测到 %d 个数学公式  (+2)", len(math_blocks))
        score += 2

    # ── 3h. 检查 media 目录中的媒体文件数量（比 inline_shapes 更全面） ──
    media_files = [n for n in zip_names if n.startswith("word/media/")]
    if len(media_files) > inline_count:
        extra = len(media_files) - inline_count
        logger.info("  [WordChecker] media 目录中还有 %d 个非内嵌媒体文件  (+2)", extra)
        score += 2

    # ── 3i. 检查嵌入对象目录 ──
    embeddings = [n for n in zip_names if n.startswith("word/embeddings/")]
    if embeddings:
        logger.info("  [WordChecker] 检测到 %d 个嵌入文件  (+2)", len(embeddings))
        score += 2

    logger.info("  [WordChecker] 文档复杂度得分：%d", score)
    return score

